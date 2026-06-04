import base64
import io
import json
from datetime import timedelta
from pathlib import Path

import qrcode
import qrcode.constants
from django.conf import settings
from django.contrib.auth import get_user_model
from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.paginator import Paginator
from django.http import HttpResponse, HttpResponseForbidden
from django.shortcuts import get_object_or_404, render
from django.urls import reverse
from django.utils import timezone
from django.views.decorators.http import require_POST
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# Vertical gap before/after agenda lines. Kept tight so a typical (trimmed)
# meeting agenda fits on a single page.
AGENDA_GAP = Pt(2)

AGENDA_TEMPLATE = (
    Path(__file__).parent / "templates" / "meetings" / "agenda" / "agenda_template.docx"
)

from .models import Meeting, MeetingRole, MeetingTypeItem, Attendance

User = get_user_model()


def _generate_qr_data_uri(url):
    """Generate a QR code as a base64-encoded PNG data URI."""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=6,
        border=2,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buffer = io.BytesIO()
    img.save(buffer, format="PNG")
    encoded = base64.b64encode(buffer.getvalue()).decode("utf-8")
    return f"data:image/png;base64,{encoded}"


def _build_agenda_sections(meeting):
    """Build an ordered list of agenda sections from the meeting's sessions.

    Returns a list of dicts:
        {"session": Session, "note": str, "roles": [MeetingRole, ...]}
    Sessions with takes_roles=False will have an empty roles list.
    Roles not assigned to any session are grouped in a final None section.
    """
    roles = list(
        meeting.roles.filter(role__show_on_agenda=True)
        .select_related("role", "user", "session", "evaluates", "evaluates__user")
        .prefetch_related("evaluators", "evaluators__user")
        .order_by("sort_order", "id")
    )

    meeting_sessions = meeting.meeting_sessions.select_related("session").order_by(
        "sort_order"
    )

    if not meeting_sessions:
        return [{"session": None, "note": "", "roles": roles}]

    sections = []
    used_role_ids = set()
    for ms in meeting_sessions:
        session = ms.session
        if session.takes_roles:
            session_roles = [r for r in roles if r.session_id == session.id]
            used_role_ids.update(r.id for r in session_roles)
        else:
            session_roles = []
        sections.append({"session": session, "note": ms.note, "roles": session_roles})

    # Roles not assigned to any session
    unassigned = [r for r in roles if r.id not in used_role_ids]
    if unassigned:
        sections.append({"session": None, "note": "", "roles": unassigned})

    return sections


def meeting_agenda(request, meeting_id):
    """Public page: presentable meeting agenda with roles and notes."""
    meeting = get_object_or_404(Meeting, id=meeting_id)
    sections = _build_agenda_sections(meeting)
    return render(
        request, "meetings/agenda.html", {"meeting": meeting, "sections": sections}
    )


def _replace_in_paragraph(paragraph, placeholder, replacement):
    """Replace placeholder text in a paragraph, handling Word's run-splitting."""
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return False
    new_text = full_text.replace(placeholder, replacement)
    # Put all text in the first run (preserves its formatting), clear the rest
    for i, run in enumerate(paragraph.runs):
        run.text = new_text if i == 0 else ""
    return True


def _replace_in_run(paragraph, placeholder, replacement):
    """Replace a placeholder that lives wholly within a single run, preserving
    that run's formatting (so e.g. a non-bold value stays non-bold next to a
    bold label). Returns False if the placeholder spans runs (caller can fall
    back to the run-collapsing _replace_in_paragraph)."""
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)
            return True
    return False


def _remove_paragraph(paragraph):
    """Remove a paragraph element from the document XML (no-op if already gone)."""
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _set_table_cell_margins(table, top_pt, bottom_pt, side_twips=108):
    """Set the table's default cell margins. Top/bottom add a little breathing
    room above and below each horizontal rule (the row borders); left/right are
    preserved near Word's default so column spacing is unchanged."""
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn("w:tblCellMar"))
    if existing is not None:
        tblPr.remove(existing)
    mar = OxmlElement("w:tblCellMar")
    for side, twips in (("top", int(top_pt * 20)), ("left", side_twips),
                        ("bottom", int(bottom_pt * 20)), ("right", side_twips)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(twips))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def _fit_table_width(table, total_twips):
    """Rescale the table's columns to span ``total_twips`` (the text-area
    width), keeping their proportions, and pin the table to a fixed layout at
    that width. The template grid was sized for the old, wider margins, so
    without this the table renders narrow and sits offset to the left."""
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    cols = grid.findall(qn("w:gridCol"))
    widths = [int(c.get(qn("w:w"))) for c in cols]
    old_total = sum(widths) or 1
    new = [round(total_twips * w / old_total) for w in widths]
    new[-1] += total_twips - sum(new)  # absorb rounding into the last column
    for col, w in zip(cols, new):
        col.set(qn("w:w"), str(w))

    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(total_twips))

    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for row in table.rows:
        for cell, w in zip(row.cells, new):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(w))


def _add_run(paragraph, text, bold=False, size=None):
    """Append a run with explicit bold/size so a line can mix bold labels with
    plain values."""
    run = paragraph.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = size
    return run


def meeting_agenda_download(request, meeting_id):
    """Public endpoint: download the meeting agenda as a Word document."""
    meeting = get_object_or_404(Meeting, id=meeting_id)

    doc = Document(AGENDA_TEMPLATE)

    # Tighten the layout so a typical (trimmed) meeting fits one page: narrow
    # margins and a 10pt body font. The template's header runs carry their own
    # sizes, so this only shrinks the role-list body text.
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.6)
    doc.styles["Normal"].font.size = Pt(10)

    # Required placeholders — always replaced
    replacements = {
        "{{DATE}}": meeting.date.strftime("%A, %B %d, %Y"),
        "{{TIME}}": meeting.date.strftime("%I:%M %p"),
    }

    # Next meeting (time, date, type) — its line is dropped if none is scheduled.
    next_meeting = (
        Meeting.objects.filter(date__gt=meeting.date).order_by("date").first()
    )
    next_meeting_text = ""
    if next_meeting:
        next_meeting_text = next_meeting.date.strftime("%I:%M %p, %A, %B %d, %Y")
        if next_meeting.meeting_type:
            next_meeting_text += f" ({next_meeting.meeting_type})"

    # Theme / Word-of-the-day share one template line. Rebuild it in code so
    # each label is bold but its value is not, and an empty field (with its
    # label) drops out. header_size is captured here and reused for the Zoom
    # link so the two match.
    header_segments = [("Theme: ", meeting.theme),
                       ("Word of the Day: ", meeting.word_of_the_day)]
    header_size = Pt(10)

    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            _replace_in_paragraph(paragraph, placeholder, value)

        text = "".join(run.text for run in paragraph.runs)

        if "{{THEME}}" in text or "{{WORD_OF_THE_DAY}}" in text:
            if paragraph.runs and paragraph.runs[0].font.size:
                header_size = paragraph.runs[0].font.size
            filled = [(label, value) for label, value in header_segments if value]
            for run in list(paragraph.runs):
                run._element.getparent().remove(run._element)
            if not filled:
                paragraphs_to_remove.append(paragraph)
            else:
                for i, (label, value) in enumerate(filled):
                    if i:
                        _add_run(paragraph, ", ", size=header_size)
                    _add_run(paragraph, label, bold=True, size=header_size)
                    _add_run(paragraph, value, size=header_size)
            continue

        # Next meeting: replace within its own (non-bold) run so the value
        # stays plain next to the bold "NEXT MEETING:" label; drop the line if
        # there is no next meeting.
        if "{{NEXT_MEETING}}" in text:
            if next_meeting_text:
                if not _replace_in_run(paragraph, "{{NEXT_MEETING}}", next_meeting_text):
                    _replace_in_paragraph(paragraph, "{{NEXT_MEETING}}", next_meeting_text)
            else:
                paragraphs_to_remove.append(paragraph)

    for p in paragraphs_to_remove:
        _remove_paragraph(p)

    # Populate the 2-column table (session | roles)
    sections = _build_agenda_sections(meeting)
    table = doc.tables[0]
    # A little breathing room above/below each horizontal rule (row border).
    _set_table_cell_margins(table, top_pt=3, bottom_pt=3)

    # Running clock: the first session starts 5 minutes after the meeting time,
    # and each session then advances by its own duration.
    current_time = meeting.date + timedelta(minutes=5)

    first = True
    for section in sections:
        session = section["session"]
        if first:
            row_cells = table.rows[0].cells
            first = False
        else:
            row_cells = table.add_row().cells

        # Cell 0: session name, then "start time, duration - note" beneath it.
        c = row_cells[0]
        p1 = c.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1.paragraph_format.space_before = AGENDA_GAP
        if session:
            r = p1.add_run(session.name)
            detail = current_time.strftime("%I:%M %p")
            if session.duration_minutes:
                detail += f", {session.duration_minutes} min"
            if section["note"]:
                detail += f" - {section['note']}"
            p2 = c.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p2.paragraph_format.space_after = AGENDA_GAP
            p2.add_run(detail)
            current_time += timedelta(minutes=session.duration_minutes or 0)
        else:
            r = p1.add_run("Other")
        r.bold = True

        # Cell 1: a main line per role (Role: Member (In Person) N min), with
        # optional italic follow-up lines for evaluator pairing and notes.
        c = row_cells[1]
        if section["roles"]:
            first_role = True
            for assignment in section["roles"]:
                if first_role:
                    p = c.paragraphs[0]
                    first_role = False
                else:
                    p = c.add_paragraph()
                p.paragraph_format.space_before = AGENDA_GAP

                member = (
                    f"{assignment.user.first_name} {assignment.user.last_name}"
                    if assignment.user
                    else "(Open)"
                )
                run = p.add_run(f"{assignment.role.name}: ")
                run.bold = True
                p.add_run(member)

                # In-person status in brackets next to the name: [L] in
                # person, [R] remote, [-] unknown.
                mode = {True: "L", False: "R"}.get(assignment.in_person, "-")
                p.add_run(f" [{mode}]")
                if assignment.time_minutes:
                    p.add_run(f" {assignment.time_minutes} min")

                # Report each evaluator pairing once, from the evaluator's
                # side ("evaluating <speaker>"); the reverse "evaluator: <name>"
                # line on the speaker's row would just duplicate it.
                for label in (
                    assignment.evaluating_label(),
                    assignment.agenda_notes(),
                ):
                    if label:
                        c.add_paragraph().add_run(label).italic = True

            c.paragraphs[-1].paragraph_format.space_after = AGENDA_GAP
        elif session and not session.takes_roles:
            p = c.paragraphs[0]
            p.paragraph_format.space_before = AGENDA_GAP
            p.paragraph_format.space_after = AGENDA_GAP
            p.add_run(section["note"] or "Break")

    # The template grid was sized for the old margins; rescale it to fill the
    # current text area so the table isn't offset to the left.
    sec = doc.sections[0]
    text_twips = round(
        (int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)) / 635
    )
    _fit_table_width(table, text_twips)

    # Final merged row welcoming members who joined within ~3 months (90 days)
    # before this meeting. Omitted entirely when there are none.
    meeting_day = meeting.date.date()
    newest = (
        get_user_model().objects.filter(
            is_active=True, is_guest=False,
            join_date__gt=meeting_day - timedelta(days=90),
            join_date__lte=meeting_day,
        )
        .order_by("join_date", "last_name", "first_name")
    )
    names = [f"{u.first_name} {u.last_name}".strip() for u in newest]
    if names:
        welcome = table.add_row()
        cell = welcome.cells[0].merge(welcome.cells[1])
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = AGENDA_GAP
        para.paragraph_format.space_after = AGENDA_GAP
        label = para.add_run("Welcome to our newest Speak Up members: ")
        label.bold = True
        para.add_run(", ".join(names))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"agenda-{meeting.date.strftime('%Y-%m-%d')}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def role_signups(request):
    """Role sign-up page: upcoming meetings with their role tables. Logged-in
    members claim or drop roles here; anonymous visitors see a read-only
    roster."""
    now = timezone.now()
    meetings_qs = (
        Meeting.objects.filter(date__gte=now)
        .order_by("date")
        .prefetch_related(
            "roles",
            "roles__role",
            "roles__user",
            "roles__evaluates__user",
            "roles__evaluators__user",
        )
    )

    paginator = Paginator(meetings_qs, 10)
    page = paginator.get_page(request.GET.get("page"))

    return render(request, "meetings/signups.html", {"meetings": page})


def _role_row_response(request, assignment, triggers=None):
    """Render the role-row partial, optionally attaching HX-Trigger events."""
    response = render(
        request, "meetings/partials/role_row.html", {"assignment": assignment}
    )
    if triggers:
        response["HX-Trigger"] = json.dumps(triggers)
    return response


def _apply_signup_fields(assignment, request):
    """Apply the member-entered fields from a sign-up or edit POST onto the
    assignment: attendance mode, notes, and — for speech roles — Pathways
    path/level/project. Shared so sign-up and edit stay consistent."""
    in_person = request.POST.get("in_person")
    if in_person == "true":
        assignment.in_person = True
    elif in_person == "false":
        assignment.in_person = False
    else:
        # The dialog requires it, but fall back to the meeting-type
        # template's expected mode if the field is somehow missing.
        assignment.in_person = _template_in_person(assignment)

    assignment.notes = request.POST.get("notes", "").strip()

    # Pathways details only apply to speech roles.
    if assignment.role.shows_pathways_fields:
        assignment.pathways_path = request.POST.get("pathways_path", "").strip()
        level = request.POST.get("pathways_level", "").strip()
        assignment.pathways_level = int(level) if level.isdigit() else None
        assignment.pathways_project = request.POST.get("pathways_project", "").strip()


def _template_in_person(assignment):
    """Default attendance mode for an assignment, read from its template.

    Looks up ``MeetingTypeItem.in_person`` for this assignment's
    (meeting_type, role). Falls back to True when the meeting has no
    ``meeting_type`` or no matching template item.
    """
    meeting_type_id = assignment.meeting.meeting_type_id
    if not meeting_type_id:
        return True
    mti = MeetingTypeItem.objects.filter(
        meeting_type_id=meeting_type_id,
        role=assignment.role,
    ).first()
    return mti.in_person if mti else True


@login_required
def signup_role_form(request, role_id):
    """HTMX endpoint: render the role dialog. For an open role it's a sign-up
    form that claims the role; for an assigned role it's an edit form (same
    fields) that updates attendance/notes/Pathways for the assignee or an
    officer."""
    assignment = get_object_or_404(MeetingRole, id=role_id)
    is_edit = assignment.user is not None
    if is_edit:
        default_in_person = (
            assignment.in_person
            if assignment.in_person is not None
            else _template_in_person(assignment)
        )
        form_action = reverse("save_role_details", args=[assignment.id])
    else:
        default_in_person = _template_in_person(assignment)
        form_action = reverse("toggle_role", args=[assignment.id])
    return render(
        request,
        "meetings/partials/signup_dialog.html",
        {
            "assignment": assignment,
            "is_edit": is_edit,
            "form_action": form_action,
            "default_in_person": default_in_person,
            "pathways_paths": MeetingRole.PATHWAYS_PATHS,
            "pathways_levels": MeetingRole.PATHWAYS_LEVELS,
        },
    )


@login_required
@require_POST
def toggle_role(request, role_id):
    """HTMX endpoint: claim or drop a meeting role. Returns the updated table row partial.

    Claiming is driven by the sign-up dialog (see ``signup_role_form``): the POST
    carries the member's attendance mode, notes, and — for speech roles — their
    Pathways path/level/project.
    """
    assignment = get_object_or_404(MeetingRole, id=role_id)

    if assignment.user == request.user:
        # Drop: user is un-signing from their own role. Clear all
        # member-entered fields so the slot reopens clean.
        assignment.user = None
        assignment.in_person = None
        assignment.notes = ""
        assignment.pathways_path = ""
        assignment.pathways_level = None
        assignment.pathways_project = ""
        assignment.save()
        return _role_row_response(request, assignment)

    elif assignment.user is None:
        # Claim: enforce one-role-per-meeting limit
        # (officers/superusers exempt), and no guest sign ups
        has_existing_role = MeetingRole.objects.filter(
            meeting=assignment.meeting, user=request.user
        ).exists()
        can_bypass = request.user.is_officer or request.user.is_superuser

        if has_existing_role and not can_bypass:
            return _role_row_response(
                request,
                assignment,
                {
                    "showAlert": "You have already signed up for a role for this meeting. Please drop your current role first.",
                    "closeModal": True,
                },
            )

        if request.user.is_guest:
            return _role_row_response(
                request,
                assignment,
                {
                    "showAlert": "Guests cannot sign up for a role. Become a member!",
                    "closeModal": True,
                },
            )

        assignment.user = request.user
        _apply_signup_fields(assignment, request)
        assignment.save()
        return _role_row_response(request, assignment, {"closeModal": True})

    else:
        # Already taken by someone else
        return HttpResponseForbidden("This role is already taken.")


@login_required
@require_POST
def save_role_details(request, role_id):
    """HTMX endpoint: edit an assigned role's attendance, notes, and Pathways
    (assignee or officer only). Uses the same fields and dialog as sign-up."""
    assignment = get_object_or_404(MeetingRole, id=role_id)

    can_edit = request.user.is_officer or (assignment.user == request.user)
    if not can_edit:
        return HttpResponseForbidden("You do not have permission to edit this role.")

    _apply_signup_fields(assignment, request)
    assignment.save()

    return _role_row_response(request, assignment, {"closeModal": True})


def checkin_kiosk(request):
    """Displays the check-in grid for today's meeting (or the next upcoming one)."""
    today = timezone.now().date()
    meeting = Meeting.objects.filter(date__date=today).first()

    if not meeting:
        meeting = Meeting.objects.filter(date__gte=today).order_by("date").first()

    context = {"meeting": meeting}

    if meeting:
        members = list(
            User.objects.filter(is_active=True)
            .exclude(username="admin")
            .order_by("first_name")
        )
        checked_in_ids = set(meeting.attendances.values_list("user_id", flat=True))

        # Map each member to their role's attendance mode for this meeting.
        # If a member holds multiple roles, the first non-null mode wins.
        mode_by_user = {}
        for uid, in_person in meeting.roles.filter(
            user__isnull=False
        ).values_list("user_id", "in_person"):
            if mode_by_user.get(uid) is None:
                mode_by_user[uid] = in_person
        for member in members:
            member.attendance_mode = mode_by_user.get(member.id)

        kiosk_url = f"{settings.SITE_URL}{reverse('checkin_kiosk')}"
        qr_data_uri = _generate_qr_data_uri(kiosk_url)
        agenda_url = reverse("meeting_agenda", args=[meeting.id])

        context.update(
            {
                "members": members,
                "checked_in_ids": checked_in_ids,
                "qr_data_uri": qr_data_uri,
                "agenda_url": agenda_url,
            }
        )

    return render(request, "meetings/kiosk.html", context)


@require_POST
def checkin_member(request, meeting_id, user_id):
    """HTMX endpoint: toggle attendance for a member (check in / undo)."""
    meeting = get_object_or_404(Meeting, id=meeting_id)
    user = get_object_or_404(User, id=user_id)

    attendance = Attendance.objects.filter(meeting=meeting, user=user).first()

    is_present = False
    if attendance:
        attendance.delete()
    else:
        Attendance.objects.create(meeting=meeting, user=user)
        is_present = True

    role = (
        meeting.roles.filter(user=user, in_person__isnull=False)
        .order_by("sort_order")
        .first()
    )
    user.attendance_mode = role.in_person if role else None

    return render(
        request,
        "meetings/partials/checkin_button.html",
        {"meeting": meeting, "member": user, "is_present": is_present},
    )


def checkin_guest(request, meeting_id):
    """POST endpoint: record a walk-in guest's name and email."""
    if request.method == "POST":
        meeting = get_object_or_404(Meeting, id=meeting_id)
        first_name = request.POST.get("guest_first_name")
        last_name = request.POST.get("guest_last_name")
        email = request.POST.get("guest_email")

        if first_name and last_name and email:
            Attendance.objects.create(
                meeting=meeting,
                guest_first_name=first_name,
                guest_last_name=last_name,
                guest_email=email,
            )
            return render(
                request, "meetings/partials/guest_success.html", {"name": first_name}
            )

    return HttpResponseForbidden()
